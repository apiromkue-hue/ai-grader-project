#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Student Project Grader Interface Page
Part of Streamlit Multi-Page App
"""

import sys
import os
import importlib.util
import time
from datetime import datetime

# Fix for Streamlit subprocess - manually load google.generativeai
def _load_google_generativeai():
    """Load google.generativeai module by finding its path directly."""
    google_genai_paths = [
        'C:\\Users\\User\\AppData\\Local\\Programs\\Python\\Python312\\lib\\site-packages\\google\\generativeai\\__init__.py',
        os.path.expanduser('~\\AppData\\Local\\Programs\\Python\\Python312\\lib\\site-packages\\google\\generativeai\\__init__.py'),
    ]
    
    for path in google_genai_paths:
        if os.path.exists(path):
            spec = importlib.util.spec_from_file_location("google.generativeai", path)
            if spec and spec.loader:
                module = importlib.util.module_from_spec(spec)
                sys.modules['google.generativeai'] = module
                spec.loader.exec_module(module)
                return module
    
    site_packages = 'C:\\Users\\User\\AppData\\Local\\Programs\\Python\\Python312\\lib\\site-packages'
    if site_packages not in sys.path:
        sys.path.insert(0, site_packages)
    
    return __import__('google.generativeai', fromlist=[''])

try:
    genai = _load_google_generativeai()
except:
    site_packages = 'C:\\Users\\User\\AppData\\Local\\Programs\\Python\\Python312\\lib\\site-packages'
    if site_packages not in sys.path:
        sys.path.insert(0, site_packages)
    import google.generativeai as genai

import streamlit as st
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# ========== PAGE CONFIG ==========
st.set_page_config(
    page_title="Student Interface - ระบบตรวจโครงงาน AI",
    page_icon="👨‍🎓",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ========== PROMPT FONT & K-MINIMAL DESIGN ==========
google_fonts = """
<link href="https://fonts.googleapis.com/css2?family=Prompt:wght@300;400;500;600;700&display=swap" rel="stylesheet">
"""
st.markdown(google_fonts, unsafe_allow_html=True)

custom_css = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Prompt:wght@300;400;500;600;700&display=swap');

* {
    font-family: 'Prompt', sans-serif !important;
}

body {
    background: linear-gradient(135deg, #F0D9E8 0%, #E8B4D4 100%);
    color: #333;
    font-family: 'Prompt', sans-serif;
}

h1, h2, h3 {
    color: #D4A5C8 !important;
    font-family: 'Prompt', sans-serif !important;
    font-weight: 600 !important;
}

h1 {
    color: #B8879F !important;
}

.stButton > button {
    background: linear-gradient(90deg, #E8B4D4 0%, #D4A5C8 100%) !important;
    color: white !important;
    border-radius: 12px !important;
    font-weight: 600 !important;
    padding: 12px 24px !important;
    transition: all 0.3s ease !important;
    border: none !important;
    font-family: 'Prompt', sans-serif !important;
    box-shadow: 0 4px 15px rgba(232, 180, 212, 0.3) !important;
}

.stButton > button:hover {
    transform: translateY(-3px) !important;
    box-shadow: 0 8px 25px rgba(212, 165, 200, 0.4) !important;
}

[data-testid="metric-container"] {
    background: linear-gradient(135deg, #FFFFFF 0%, #F5E8F0 100%);
    border-radius: 15px;
    padding: 20px;
    box-shadow: 0 4px 15px rgba(232, 180, 212, 0.2);
    border: 2px solid #F0D9E8;
}

[data-testid="metric-container"]:hover {
    transform: translateY(-5px);
    box-shadow: 0 8px 25px rgba(212, 165, 200, 0.3);
    border-color: #E8B4D4;
}

hr {
    border: 0;
    height: 2px;
    background: linear-gradient(90deg, #E8B4D4 0%, #D4A5C8 100%);
    margin: 20px 0;
}

/* ปรับ padding ด้านบนให้พอดี */
.main .block-container {
    padding-top: 1rem !important;
}

</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ========== SIDEBAR REMOVED ==========

# ========== INITIALIZATION ==========
if 'student_logged_in' not in st.session_state:
    st.session_state.student_logged_in = False
if 'student_username' not in st.session_state:
    st.session_state.student_username = None

# ========== LOGIN PAGE ==========
if not st.session_state.student_logged_in:
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        st.markdown("# 👨‍🎓 ระบบตรวจโครงงาน AI")
        st.markdown("### Student Interface")
        
        with st.form("student_login_form"):
            st.markdown("#### เข้าสู่ระบบ")
            username = st.text_input("👤 Username")
            password = st.text_input("🔐 Password", type="password")
            
            submitted = st.form_submit_button("🚀 เข้าสู่ระบบ", use_container_width=True)
            
            if submitted:
                if username and password:
                    # ตรวจสอบจาก users_database.json
                    import json
                    try:
                        users_db_file = "users_database.json"
                        if os.path.exists(users_db_file):
                            with open(users_db_file, 'r', encoding='utf-8') as f:
                                users_db = json.load(f)
                            
                            # ค้นหา user
                            found_user = None
                            if "users" in users_db:
                                for user in users_db["users"]:
                                    if user.get("username") == username and \
                                       user.get("password") == password and \
                                       user.get("status") == "active":
                                        found_user = user
                                        break
                            
                            if found_user:
                                st.session_state.student_logged_in = True
                                st.session_state.student_username = username
                                st.session_state.student_name = found_user.get("name", username)
                                st.session_state.student_role = found_user.get("role", "student")
                                
                                # อัพเดท last_login
                                found_user["last_login"] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                                with open(users_db_file, 'w', encoding='utf-8') as f:
                                    json.dump(users_db, f, ensure_ascii=False, indent=2)
                                
                                st.success("✅ เข้าสู่ระบบสำเร็จ!")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.error("❌ Username หรือ Password ไม่ถูกต้อง")
                        else:
                            st.error("❌ ไม่พบฐานข้อมูลผู้ใช้")
                    except Exception as e:
                        st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
                else:
                    st.warning("⚠️ กรุณากรอก Username และ Password")

# ========== MAIN INTERFACE ==========
else:
    # หัวข้อหน้า
    st.markdown("# 👨‍🎓 ส่วนของนักเรียนและอาจารย์")
    st.markdown(f"ยินดีต้อนรับ **{st.session_state.student_username}**! 👋")
    
    st.markdown("---")
    
    tab1, tab2, tab3, tab4 = st.tabs(["📂 อัปโหลดและวิเคราะห์", "📜 ประวัติการวิเคราะห์", "📊 สถิติและรายงาน", "🚪 ออกจากระบบ"])
    
    with tab1:
        st.markdown("## 📂 อัปโหลดไฟล์โครงงาน")
        
        # ตรวจสอบ API Key จากหลายแหล่ง
        def get_api_key():
            # 1. จาก Environment Variable (.env file) - รองรับทั้ง 2 ชื่อ
            env_key = os.environ.get("GOOGLE_API_KEY") or os.environ.get("GOOGLE_GEMINI_API_KEY")
            if env_key and env_key != "your_api_key_here":
                return env_key
            
            # 2. จาก Streamlit Secrets (สำหรับ deployment)
            try:
                if hasattr(st, 'secrets'):
                    return st.secrets.get('GOOGLE_API_KEY') or st.secrets.get('GOOGLE_GEMINI_API_KEY')
            except:
                pass
            
            # 3. จาก Session State
            if 'gemini_api_key' in st.session_state:
                return st.session_state.gemini_api_key
            
            return None
        
        # ตรวจสอบ API Key แต่ไม่แสดง
        current_api_key = get_api_key()
        if not current_api_key:
            st.warning("⚠️ ยังไม่มี API Key กรุณาตั้งค่า")
            
            with st.expander("🔑 ตั้งค่า Google Gemini API Key", expanded=True):
                st.markdown("""
                ### วิธีการตั้งค่า API Key (เลือก 1 วิธี):
                
                #### 🔒 วิธีที่ 1: ใช้ไฟล์ .env (ปลอดภัยที่สุด - แนะนำ)
                1. สร้างไฟล์ `.env` ในโฟลเดอร์โปรเจค
                2. เพิ่มบรรทัด: `GOOGLE_GEMINI_API_KEY=your_api_key_here`
                3. รีสตาร์ท Streamlit
                
                #### 🌐 วิธีที่ 2: ใช้ Streamlit Secrets (สำหรับ Cloud)
                1. สร้างโฟลเดอร์ `.streamlit/`
                2. สร้างไฟล์ `.streamlit/secrets.toml`
                3. เพิ่ม: `GOOGLE_GEMINI_API_KEY = "your_api_key_here"`
                
                #### 💻 วิธีที่ 3: ใส่ชั่วคราว (session นี้เท่านั้น)
                """)
                
                st.info("💡 ขอ API Key ฟรีที่: [Google AI Studio](https://makersuite.google.com/app/apikey)")
                
                api_key_input = st.text_input(
                    "ใส่ API Key ชั่วคราว",
                    type="password",
                    placeholder="AIzaSy...",
                    help="API Key นี้จะหายเมื่อปิดเบราว์เซอร์"
                )
                
                if api_key_input:
                    st.session_state.gemini_api_key = api_key_input
                    st.success("✅ บันทึก API Key ชั่วคราวแล้ว")
                    st.rerun()
        
        uploaded_file = st.file_uploader("เลือกไฟล์ PDF หรือ Word", type=["pdf", "docx"])
        
        if uploaded_file:
            col1, col2 = st.columns([1, 1])
            with col1:
                st.info(f"📄 File: {uploaded_file.name}")
            with col2:
                st.info(f"📊 Size: {uploaded_file.size / 1024:.2f} KB")
            
            st.markdown("---")
            
            # เลือกบทที่ต้องการตรวจ
            st.markdown("### 📚 เลือกบทที่ต้องการตรวจ")
            
            chapter_options = {
                "ทั้งหมด (5 บท)": "all",
                "บทที่ 1: บทนำ": "chapter1",
                "บทที่ 2: เอกสารและงานวิจัยที่เกี่ยวข้อง": "chapter2",
                "บทที่ 3: วิธีดำเนินการวิจัย": "chapter3",
                "บทที่ 4: ผลการวิจัย": "chapter4",
                "บทที่ 5: สรุปและข้อเสนอแนะ": "chapter5"
            }
            
            selected_chapter = st.selectbox(
                "เลือกบทที่ต้องการวิเคราะห์",
                options=list(chapter_options.keys()),
                index=0,
                help="เลือกบทเฉพาะเพื่อตรวจละเอียด หรือเลือกทั้งหมดเพื่อตรวจทั้งโครงงาน"
            )
            
            chapter_value = chapter_options[selected_chapter]
            
            # แสดงสถานะความคืบหน้าของแต่ละบท (ถ้ามีข้อมูลเก่า)
            if 'chapter_progress' not in st.session_state:
                st.session_state.chapter_progress = {}
            
            user_key = f"{st.session_state.student_username}_{uploaded_file.name}"
            if user_key in st.session_state.chapter_progress:
                st.markdown("#### 📊 ความคืบหน้าการตรวจโครงงาน")
                progress_data = st.session_state.chapter_progress[user_key]
                
                cols = st.columns(5)
                for i, (chapter_name, status) in enumerate([
                    ("บทที่ 1", progress_data.get('chapter1', False)),
                    ("บทที่ 2", progress_data.get('chapter2', False)),
                    ("บทที่ 3", progress_data.get('chapter3', False)),
                    ("บทที่ 4", progress_data.get('chapter4', False)),
                    ("บทที่ 5", progress_data.get('chapter5', False))
                ]):
                    with cols[i]:
                        if status:
                            st.success(f"✅ {chapter_name}")
                        else:
                            st.info(f"⏳ {chapter_name}")
                
                completed = sum(1 for v in progress_data.values() if v)
                total = 5
                progress_percent = (completed / total) * 100
                st.progress(progress_percent / 100)
                st.markdown(f"**ความคืบหน้าทั้งหมด:** {completed}/{total} บท ({progress_percent:.0f}%)")
                st.markdown("---")
            
            if st.button("🚀 เริ่มวิเคราะห์", type="primary", use_container_width=True):
                try:
                    # อ่านเนื้อหาไฟล์จริง
                    with st.spinner("📄 กำลังอ่านไฟล์..."):
                        file_content = ""
                        num_pages = 0
                        
                        if uploaded_file.type == "application/pdf":
                            try:
                                import PyPDF2
                                uploaded_file.seek(0)
                                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                                num_pages = len(pdf_reader.pages)
                                
                                for page_num, page in enumerate(pdf_reader.pages):
                                    text = page.extract_text()
                                    file_content += f"\n=== หน้า {page_num + 1} ===\n{text}\n"
                                
                                if not file_content.strip():
                                    st.error("❌ ไม่สามารถอ่านข้อความจาก PDF ได้")
                                    st.stop()
                            except Exception as e:
                                st.error(f"❌ เกิดข้อผิดพลาดในการอ่านไฟล์ PDF: {str(e)}")
                                st.stop()
                        
                        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            try:
                                from docx import Document
                                uploaded_file.seek(0)
                                doc = Document(uploaded_file)
                                num_pages = len(doc.paragraphs) // 30 + 1  # ประมาณการ
                                
                                for para in doc.paragraphs:
                                    file_content += para.text + "\n"
                                
                                if not file_content.strip():
                                    st.error("❌ ไม่สามารถอ่านข้อความจาก Word ได้")
                                    st.stop()
                            except Exception as e:
                                st.error(f"❌ เกิดข้อผิดพลาดในการอ่านไฟล์ Word: {str(e)}")
                                st.stop()
                        
                        else:
                            st.error("❌ รองรับเฉพาะไฟล์ PDF และ Word เท่านั้น")
                            st.stop()
                        
                        # ตรวจสอบความยาว
                        word_count = len(file_content.split())
                        char_count = len(file_content)
                        
                        st.success(f"✅ อ่านไฟล์สำเร็จ: {num_pages} หน้า, {word_count:,} คำ, {char_count:,} ตัวอักษร")
                    
                    # ตรวจสอบ API Key
                    api_key = get_api_key()
                    if not api_key:
                        st.error("❌ กรุณาใส่ Google Gemini API Key ก่อนวิเคราะห์")
                        st.info("💡 กลับไปด้านบนเพื่อตั้งค่า API Key")
                        st.stop()
                    
                    # ใช้ AI วิเคราะห์จริง
                    chapter_focus = ""
                    if chapter_value != "all":
                        chapter_names = {
                            "chapter1": "บทที่ 1: บทนำ",
                            "chapter2": "บทที่ 2: เอกสารและงานวิจัยที่เกี่ยวข้อง",
                            "chapter3": "บทที่ 3: วิธีดำเนินการวิจัย",
                            "chapter4": "บทที่ 4: ผลการวิจัย",
                            "chapter5": "บทที่ 5: สรุปและข้อเสนอแนะ"
                        }
                        chapter_focus = f"\n\n**⚠️ หมายเหตุ:** กรุณาวิเคราะห์เฉพาะ {chapter_names.get(chapter_value, '')} เท่านั้น\n"
                    
                    with st.spinner(f"🤖 กำลังใช้ AI วิเคราะห์{' ' + selected_chapter if chapter_value != 'all' else 'ทั้งหมด'}..."):
                        # ใช้ SDK
                        genai.configure(api_key=api_key)
                        model = genai.GenerativeModel('gemini-2.5-flash')
                                
                        # สร้าง Prompt สำหรับ Semantic Analysis 3 ระดับ
                        analysis_prompt = f"""
คุณคือผู้เชี่ยวชาญด้านการตรวจสอบและประเมินโครงงานวิทยาศาสตร์ ให้วิเคราะห์เอกสารโครงงานนี้แบบ Semantic Analysis 3 ระดับ:

**ข้อมูลไฟล์:**
- ชื่อไฟล์: {uploaded_file.name}
- จำนวนหน้า: {num_pages} หน้า
- จำนวนคำ: {word_count:,} คำ
- **บทที่ต้องการตรวจ:** {selected_chapter}
{chapter_focus}

**เนื้อหาเอกสาร:**
{file_content[:15000]}

{"... (เนื้อหาถูกตัดเนื่องจากความยาว)" if len(file_content) > 15000 else ""}

---

**กรุณาวิเคราะห์ตามโครงสร้างนี้เท่านั้น:**

## ระดับที่ 1: การตรวจรูปแบบ (Format Check) - คะแนนเต็ม 30
วิเคราะห์{"เฉพาะ" + selected_chapter if chapter_value != "all" else ""}:
- มีโครงสร้างครบถ้วนหรือไม่?
- มีบรรณานุกรม/เอกสารอ้างอิงหรือไม่? (ถ้าเป็นบททั้งหมด)
- การจัดรูปแบบเหมาะสมหรือไม่?
- ให้คะแนนและระบุข้อบกพร่อง

## ระดับที่ 2: การตรวจความเชื่อมโยง (Logical Consistency) - คะแนนเต็ม 40
วิเคราะห์{"เฉพาะ" + selected_chapter if chapter_value != "all" else ""}:
1. {"วัตถุประสงค์สอดคล้องกับเนื้อหาหรือไม่?" if chapter_value == "chapter1" else "เนื้อหามีความสอดคล้องภายในหรือไม่?"}
2. ข้อมูลที่ระบุมีความสมเหตุสมผลหรือไม่?
3. {"ทฤษฎีเชื่อมโยงกับการวิจัยหรือไม่?" if chapter_value == "chapter2" else "มีความเชื่อมโยงระหว่างส่วนต่างๆ หรือไม่?"}
- ให้คะแนนและระบุความไม่สอดคล้อง

## ระดับที่ 3: ข้อเสนอแนะเชิงเนื้อหา (Content Feedback) - คะแนนเต็ม 30
ให้คำแนะนำเฉพาะเจาะจง:
1. บทคัดย่อยาวเกินไป (>300 คำ) หรือไม่?
2. วัตถุประสงค์เป็นข้อ ๆ ชัดเจนหรือไม่?
3. มีการอ้างอิงครบถ้วนหรือไม่?
4. ควรปรับปรุงอะไรบ้าง?

**สรุปคะแนนรวม: X/100**
- ระดับ 1: X/30
- ระดับ 2: X/40  
- ระดับ 3: X/30

**แผนการปรับปรุง 3 ข้อสำคัญที่สุด**

กรุณาวิเคราะห์อย่างละเอียดและให้คำแนะนำที่เป็นประโยชน์จริง ๆ
"""
                        
                        # ส่งไปให้ AI วิเคราะห์
                        response = model.generate_content(analysis_prompt)
                        ai_analysis = response.text
                        ai_model_used = "Google Gemini Pro (Real AI)"
                    
                    # สร้างรายงานฉบับสมบูรณ์
                    analysis_result = f"""
# 📊 ผลการวิเคราะห์โครงงานแบบ Semantic Analysis

**📄 ชื่อไฟล์:** {uploaded_file.name}  
**📏 ขนาดไฟล์:** {uploaded_file.size / 1024:.2f} KB  
**📑 จำนวนหน้า:** {num_pages} หน้า  
**📝 จำนวนคำ:** {word_count:,} คำ  
**🔤 จำนวนตัวอักษร:** {char_count:,} ตัวอักษร

---

## 🤖 ผลการวิเคราะห์

{ai_analysis}

---

**📅 วันที่วิเคราะห์:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}  
**👤 ผู้วิเคราะห์:** {st.session_state.student_username}  
**🤖 AI Model:** {ai_model_used}  
**⚡ ระดับการวิเคราะห์:** 3 ระดับ (Format + Consistency + Content)

---

💡 **หมายเหตุ:** ผลการวิเคราะห์นี้เป็นการประเมินอัตโนมัติ ควรตรวจสอบโดยอาจารย์ที่ปรึกษาอีกครั้ง
"""
                    
                    # บันทึกผลการวิเคราะห์ลงฐานข้อมูล
                    try:
                        import json
                        
                        # อ่านฐานข้อมูล
                        history_file = "history.json"
                        if os.path.exists(history_file):
                            with open(history_file, 'r', encoding='utf-8') as f:
                                history_data = json.load(f)
                        else:
                            history_data = {}
                        
                        # ตรวจสอบว่ามี key "analyses" หรือไม่
                        if "analyses" not in history_data:
                            history_data["analyses"] = []
                        
                        # อัพเดทความคืบหน้าของบทที่ตรวจ
                        if user_key not in st.session_state.chapter_progress:
                            st.session_state.chapter_progress[user_key] = {}
                        
                        if chapter_value == "all":
                            # ถ้าตรวจทั้งหมด = ทุกบทเสร็จ
                            st.session_state.chapter_progress[user_key] = {
                                'chapter1': True,
                                'chapter2': True,
                                'chapter3': True,
                                'chapter4': True,
                                'chapter5': True
                            }
                        else:
                            # เฉพาะบทที่เลือก
                            st.session_state.chapter_progress[user_key][chapter_value] = True
                        
                        # แยกคะแนนจากผลการวิเคราะห์
                        score = None
                        try:
                            import re
                            # หาคะแนนรวมจากรูปแบบ "สรุปคะแนนรวม: XX/100"
                            score_match = re.search(r'สรุปคะแนนรวม[:\s]*([0-9]+)/100', ai_analysis)
                            if score_match:
                                score = int(score_match.group(1))
                            else:
                                # ลองหารูปแบบ "คะแนน: XX/100"
                                score_match = re.search(r'คะแนน[:\s]*([0-9]+)/100', ai_analysis)
                                if score_match:
                                    score = int(score_match.group(1))
                        except:
                            score = None
                        
                        # เพิ่มข้อมูลใหม่
                        new_entry = {
                            "id": f"ANALYSIS_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                            "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "username": st.session_state.student_username,
                            "file_name": uploaded_file.name,
                            "chapter_checked": selected_chapter,
                            "chapter_value": chapter_value,
                            "file_size": uploaded_file.size,
                            "num_pages": num_pages,
                            "word_count": word_count,
                            "char_count": char_count,
                            "score": score,
                            "analysis_result": ai_analysis,
                            "full_report": analysis_result,
                            "chapter_progress": st.session_state.chapter_progress[user_key]
                        }
                        
                        history_data["analyses"].append(new_entry)
                        
                        # บันทึกกลับ
                        with open(history_file, 'w', encoding='utf-8') as f:
                            json.dump(history_data, f, ensure_ascii=False, indent=2)
                        
                    except Exception as e:
                        st.warning(f"⚠️ ไม่สามารถบันทึกลงฐานข้อมูลได้: {str(e)}")
                    
                    st.success("✅ วิเคราะห์เสร็จสมบูรณ์!")
                    st.markdown(analysis_result)
                
                except Exception as e:
                    error_msg = str(e)
                    st.error(f"❌ เกิดข้อผิดพลาดในการวิเคราะห์: {error_msg}")
                    
                    if "API_KEY_INVALID" in error_msg or "API key not valid" in error_msg:
                        st.error("🔑 API Key ไม่ถูกต้อง กรุณาตรวจสอบและใส่ใหม่")
                        st.info("💡 ไปที่ [Google AI Studio](https://makersuite.google.com/app/apikey) เพื่อสร้าง API Key ใหม่")
                    elif "RESOURCE_EXHAUSTED" in error_msg or "quota" in error_msg.lower():
                        st.error("⚠️ API Key หมดโควต้าแล้ว กรุณาใช้ API Key ใหม่หรือรอ 24 ชั่วโมง")
                    else:
                        st.error("กรุณาตรวจสอบ Google API Key หรือลองใหม่อีกครั้ง")
                    
                    # ลบ API Key เก่าออก
                    if 'gemini_api_key' in st.session_state:
                        del st.session_state.gemini_api_key
                        
                    st.stop()
                    
                # ปุ่มดาวน์โหลดรายงาน (inside try block)
                if 'analysis_result' in locals():
                    col_dl1, col_dl2 = st.columns(2)
                    with col_dl1:
                        st.download_button(
                        "📥 ดาวน์โหลดรายงาน PDF",
                        data=analysis_result,
                        file_name=f"analysis_{uploaded_file.name}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                    with col_dl2:
                        st.download_button(
                            "📄 ดาวน์โหลดรายงาน Word",
                            data=analysis_result,
                            file_name=f"analysis_{uploaded_file.name}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
    
    with tab2:
        st.markdown("## 📜 ประวัติการวิเคราะห์")
        
        # อ่านประวัติจาก history.json
        try:
            import json
            history_file = "history.json"
            
            if os.path.exists(history_file):
                with open(history_file, 'r', encoding='utf-8') as f:
                    history_data = json.load(f)
                
                # กรองเฉพาะของ user ที่ login
                if "analyses" in history_data:
                    user_analyses = [
                        entry for entry in history_data["analyses"]
                        if entry.get("username") == st.session_state.student_username
                    ]
                    
                    if user_analyses:
                        # เรียงจากใหม่ไปเก่า
                        user_analyses.sort(key=lambda x: x.get("timestamp", ""), reverse=True)
                        
                        # แสดงเป็นตาราง
                        display_data = {
                            "ไฟล์": [entry.get("file_name", "-") for entry in user_analyses],
                            "บทที่ตรวจ": [entry.get("chapter_checked", "ทั้งหมด") for entry in user_analyses],
                            "วันที่": [entry.get("timestamp", "-") for entry in user_analyses],
                            "จำนวนคำ": [f"{entry.get('word_count', 0):,}" for entry in user_analyses],
                            "สถานะ": ["✅ สำเร็จ" for _ in user_analyses]
                        }
                        
                        st.dataframe(display_data, use_container_width=True)
                        
                        # แสดงรายละเอียดเพิ่มเติม
                        st.markdown("---")
                        st.markdown("### 📊 รายละเอียดการวิเคราะห์แต่ละครั้ง")
                        
                        for i, entry in enumerate(user_analyses[:5], 1):  # แสดงแค่ 5 รายการล่าสุด
                            with st.expander(f"📄 {entry.get('file_name', 'ไม่ระบุ')} - {entry.get('timestamp', '-')}"):
                                col_a, col_b = st.columns(2)
                                with col_a:
                                    st.markdown(f"**บทที่ตรวจ:** {entry.get('chapter_checked', 'ทั้งหมด')}")
                                    st.markdown(f"**จำนวนคำ:** {entry.get('word_count', 0):,} คำ")
                                with col_b:
                                    st.markdown(f"**จำนวนหน้า:** {entry.get('num_pages', 0)} หน้า")
                                    st.markdown(f"**ขนาดไฟล์:** {entry.get('file_size', 0) / 1024:.2f} KB")
                                
                                # แสดงความคืบหน้า
                                if 'chapter_progress' in entry:
                                    progress = entry['chapter_progress']
                                    completed = sum(1 for v in progress.values() if v)
                                    total = 5
                                    st.progress(completed / total)
                                    st.markdown(f"**ความคืบหน้า:** {completed}/{total} บท ({completed/total*100:.0f}%)")
                                
                                # แสดงผลการวิเคราะห์
                                if st.button(f"📖 ดูผลการวิเคราะห์", key=f"view_{i}"):
                                    st.markdown(entry.get('analysis_result', 'ไม่มีข้อมูล'))
                    else:
                        st.info("ℹ️ ยังไม่มีประวัติการวิเคราะห์")
                else:
                    st.info("ℹ️ ยังไม่มีประวัติการวิเคราะห์")
            else:
                st.info("ℹ️ ยังไม่มีประวัติการวิเคราะห์")
                
        except Exception as e:
            st.error(f"❌ ไม่สามารถโหลดประวัติได้: {str(e)}")
            st.info("ℹ️ กรุณาวิเคราะห์ไฟล์เพื่อสร้างประวัติใหม่")
    
    with tab3:
        st.markdown("## 📊 สถิติของคุณ")
        
        # อ่านข้อมูลจริงจาก history.json
        try:
            import json
            history_file = "history.json"
            
            total_files = 0
            total_words = 0
            total_pages = 0
            total_score = 0
            score_count = 0
            chapter_stats = {}
            
            if os.path.exists(history_file):
                with open(history_file, 'r', encoding='utf-8') as f:
                    history_data = json.load(f)
                
                # กรองเฉพาะของ user ที่ login
                if "analyses" in history_data:
                    user_analyses = [
                        entry for entry in history_data["analyses"]
                        if entry.get("username") == st.session_state.student_username
                    ]
                    
                    # คำนวณสถิติ
                    total_files = len(user_analyses)
                    
                    for entry in user_analyses:
                        total_words += entry.get('word_count', 0)
                        total_pages += entry.get('num_pages', 0)
                        
                        # นับคะแนน
                        if entry.get('score') is not None:
                            total_score += entry.get('score', 0)
                            score_count += 1
                        
                        # นับบทที่ตรวจ
                        chapter = entry.get('chapter_checked', 'ไม่ระบุ')
                        chapter_stats[chapter] = chapter_stats.get(chapter, 0) + 1
            
            # แสดงสถิติ
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📂 ไฟล์ที่วิเคราะห์", f"{total_files}", 
                         delta=None if total_files == 0 else f"+{total_files}")
            with col2:
                avg_words = total_words // total_files if total_files > 0 else 0
                st.metric("📝 จำนวนคำเฉลี่ย", f"{avg_words:,}", 
                         delta=None)
            with col3:
                avg_score = total_score // score_count if score_count > 0 else 0
                score_display = f"{avg_score}/100" if score_count > 0 else "ยังไม่มีข้อมูล"
                st.metric("⭐ คะแนนเฉลี่ย", score_display, 
                         delta=None if score_count == 0 else ("+5" if avg_score >= 80 else "ปรับปรุง"))
            
            # สถิติเพิ่มเติม
            if total_files > 0:
                st.markdown("---")
                st.markdown("### 📈 สถิติเพิ่มเติม")
                
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("📄 จำนวนหน้ารวม", f"{total_pages:,}")
                with col_b:
                    st.metric("📊 จำนวนคำรวม", f"{total_words:,}")
                with col_c:
                    avg_pages = total_pages / total_files if total_files > 0 else 0
                    st.metric("📃 หน้าเฉลี่ยต่อไฟล์", f"{avg_pages:.1f}")
                
                # กราฟแท่งแสดงการตรวจแต่ละบท
                if chapter_stats:
                    st.markdown("---")
                    st.markdown("### 📊 สถิติการตรวจแต่ละบท")
                    
                    import pandas as pd
                    df_chapters = pd.DataFrame({
                        'บท': list(chapter_stats.keys()),
                        'จำนวนครั้ง': list(chapter_stats.values())
                    })
                    st.bar_chart(df_chapters.set_index('บท'))
                
                # แสดงประวัติคะแนน (ถ้ามี)
                if score_count > 0:
                    st.markdown("---")
                    st.markdown("### 📈 ประวัติคะแนน")
                    
                    scores_data = [
                        entry.get('score', 0) for entry in user_analyses 
                        if entry.get('score') is not None
                    ]
                    
                    if scores_data:
                        import pandas as pd
                        df_scores = pd.DataFrame({
                            'ลำดับ': [f"ครั้งที่ {i+1}" for i in range(len(scores_data))],
                            'คะแนน': scores_data
                        })
                        st.line_chart(df_scores.set_index('ลำดับ'))
            else:
                st.info("ℹ️ ยังไม่มีข้อมูลสถิติ กรุณาวิเคราะห์ไฟล์ก่อน")
                
        except Exception as e:
            st.error(f"❌ ไม่สามารถโหลดสถิติได้: {str(e)}")
            st.info("ℹ️ กรุณาวิเคราะห์ไฟล์เพื่อสร้างสถิติใหม่")
    
    with tab4:
        st.markdown("## 🚪 ออกจากระบบ")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("""<div style="
                background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
                border: 2px solid #ffc107;
                border-radius: 15px;
                padding: 2rem;
                text-align: center;
                margin: 2rem 0;
            ">
                <h3 style="color: #856404; margin: 0;">⚠️ คุณต้องการออกจากระบบหรือไม่?</h3>
            </div>""", unsafe_allow_html=True)
            
            st.markdown("""<div style="text-align: center; margin: 1rem 0;">
                <p style="font-size: 1.1rem; color: #666;">
                    หากคุณออกจากระบบ คุณจะต้องเข้าสู่ระบบอีกครั้ง<br>
                    เพื่อใช้งานฟีเจอร์ต่างๆ ของระบบ
                </p>
            </div>""", unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            col_a, col_b = st.columns(2)
            
            with col_a:
                if st.button("🚪 ยืนยันออกจากระบบ", type="primary", use_container_width=True):
                    # ล้างข้อมูล session
                    st.session_state.student_logged_in = False
                    st.session_state.student_username = None
                    st.session_state.student_name = None
                    st.session_state.student_role = None
                    st.session_state.logged_in = False
                    st.session_state.username = None
                    
                    # ลบข้อมูลอื่นๆ
                    if 'chapter_progress' in st.session_state:
                        del st.session_state.chapter_progress
                    if 'gemini_api_key' in st.session_state:
                        del st.session_state.gemini_api_key
                    
                    st.success("✅ ออกจากระบบสำเร็จ!")
                    time.sleep(1)
                    
                    # กลับไปหน้าหลัก
                    st.markdown("""
                    <meta http-equiv="refresh" content="0; url=/" />
                    <script>
                        window.parent.location.href = '/';
                    </script>
                    """, unsafe_allow_html=True)
                    st.stop()
            
            with col_b:
                st.markdown("""
                <a href="/" target="_self">
                    <button style="
                        width: 100%;
                        padding: 0.5rem 1rem;
                        background: #6c757d;
                        color: white;
                        border: none;
                        border-radius: 12px;
                        font-family: 'Prompt', sans-serif;
                        font-size: 16px;
                        font-weight: 600;
                        cursor: pointer;
                        box-shadow: 0 4px 15px rgba(108, 117, 125, 0.3);
                    ">
                        ❌ ยกเลิก
                    </button>
                </a>
                """, unsafe_allow_html=True)

