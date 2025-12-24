#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Report generator module for exporting analysis results to Word and PDF formats
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Preformatted
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY


class ReportGenerator:
    """สร้างรายงานการวิเคราะห์ในรูปแบบต่างๆ"""
    
    @staticmethod
    def generate_word_report(
        username: str,
        file_name: str,
        analysis_result: str,
        timestamp: str = None
    ) -> io.BytesIO:
        """
        สร้างรายงาน Word Document
        
        Args:
            username: ชื่อผู้ใช้
            file_name: ชื่อไฟล์ที่วิเคราะห์
            analysis_result: ผลการวิเคราะห์จาก AI
            timestamp: วันเวลา (ถ้าไม่มีจะใช้เวลาปัจจุบัน)
        
        Returns:
            BytesIO object ของไฟล์ Word
        """
        if timestamp is None:
            timestamp = datetime.now().isoformat()
        
        # สร้าง Document
        doc = Document()
        
        # ตั้งค่า margin
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header
        header = doc.add_heading("รายงานการวิเคราะห์โครงงาน", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_format = header.runs[0]
        header_format.font.color.rgb = RGBColor(46, 134, 222)  # Blue color
        
        # Subtitle
        subtitle = doc.add_heading("ระบบตรวจความสอดคล้องโครงงาน AI", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Metadata Section
        doc.add_heading("📋 ข้อมูลการวิเคราะห์", level=2)
        
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        # Set column widths
        metadata_table.autofit = False
        metadata_table.allow_autofit = False
        
        # Fill metadata
        metadata = [
            ("ชื่อผู้ใช้", username),
            ("ชื่อไฟล์", file_name),
            ("วันเวลาวิเคราะห์", timestamp),
            ("ระบบ", "AI Project Grader (Google Gemini)")
        ]
        
        for idx, (key, value) in enumerate(metadata):
            cells = metadata_table.rows[idx].cells
            cells[0].text = key
            cells[1].text = value
            
            # Style key column
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Analysis Result Section
        doc.add_heading("📊 ผลการวิเคราะห์", level=2)
        
        # Parse markdown-like text to add formatting
        result_lines = analysis_result.split('\n')
        for line in result_lines:
            if line.startswith('##'):
                # Heading
                heading_text = line.replace('##', '').strip()
                doc.add_heading(heading_text, level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                text = line.replace('**', '')
                p = doc.add_paragraph(text)
                for run in p.runs:
                    run.font.bold = True
            elif line.startswith('-'):
                # Bullet point
                text = line.replace('-', '').strip()
                doc.add_paragraph(text, style='List Bullet')
            elif line.strip():
                # Regular text
                doc.add_paragraph(line)
        
        doc.add_paragraph()  # Spacing
        
        # Footer
        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            f"สร้างโดย: AI Project Grader System | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
    
    @staticmethod
    def get_word_filename(file_name: str, username: str) -> str:
        """
        สร้างชื่อไฟล์ Word report
        
        Args:
            file_name: ชื่อไฟล์ต้นฉบับ
            username: ชื่อผู้ใช้
        
        Returns:
            ชื่อไฟล์ Word
        """
        # ลบ extension จากชื่อไฟล์เดิม
        original_name = file_name.split('.')[0]
        
        # สร้างชื่อใหม่
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = f"รายงาน_{original_name}_{username}_{timestamp}.docx"
        
        return report_name
    
    @staticmethod
    def generate_pdf_report(
        username: str,
        file_name: str,
        analysis_result: str,
        timestamp: str = None
    ) -> io.BytesIO:
        """
        สร้างรายงาน PDF
        
        Args:
            username: ชื่อผู้ใช้
            file_name: ชื่อไฟล์ที่วิเคราะห์
            analysis_result: ผลการวิเคราะห์จาก AI
            timestamp: วันเวลา
        
        Returns:
            BytesIO object ของไฟล์ PDF
        """
        if timestamp is None:
            timestamp = datetime.now().isoformat()
        
        # สร้าง BytesIO object สำหรับ PDF
        pdf_buffer = io.BytesIO()
        
        # สร้าง PDF document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=1*inch,
            bottomMargin=0.75*inch
        )
        
        # สร้าง styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        )
        
        # เก็บเนื้อหา
        story = []
        
        # Header
        story.append(Paragraph("รายงานการวิเคราะห์โครงงาน", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata table
        metadata = [
            [Paragraph("<b>ชื่อผู้ใช้:</b>", normal_style), Paragraph(username, normal_style)],
            [Paragraph("<b>ชื่อไฟล์:</b>", normal_style), Paragraph(file_name, normal_style)],
            [Paragraph("<b>วันเวลา:</b>", normal_style), Paragraph(timestamp, normal_style)],
            [Paragraph("<b>ระบบ:</b>", normal_style), Paragraph("AI Project Grader", normal_style)]
        ]
        
        metadata_table = Table(metadata, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f4ff')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')])
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Analysis result section
        story.append(Paragraph("📊 ผลการวิเคราะห์", heading_style))
        story.append(Spacer(1, 0.1*inch))
        
        # Parse markdown-like text
        result_lines = analysis_result.split('\n')
        for line in result_lines:
            if line.startswith('##'):
                # Heading
                heading_text = line.replace('##', '').strip()
                story.append(Paragraph(heading_text, heading_style))
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                text = line.replace('**', '')
                story.append(Paragraph(f"<b>{text}</b>", normal_style))
            elif line.startswith('-'):
                # Bullet point
                text = line.replace('-', '').strip()
                story.append(Paragraph(f"• {text}", normal_style))
            elif line.strip():
                # Regular text
                story.append(Paragraph(line, normal_style))
            else:
                story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Footer
        footer_text = f"สร้างโดย: AI Project Grader System | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(footer_text, footer_style))
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    
    @staticmethod
    def get_pdf_filename(file_name: str, username: str) -> str:
        """สร้างชื่อไฟล์ PDF"""
        original_name = file_name.split('.')[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"รายงาน_{original_name}_{username}_{timestamp}.pdf"
    
    @staticmethod
    def generate_summary_report(username: str, history_list: list) -> io.BytesIO:
        """
        สร้างรายงานสรุปประวัติการวิเคราะห์ทั้งหมด
        
        Args:
            username: ชื่อผู้ใช้
            history_list: รายการประวัติการวิเคราะห์
        
        Returns:
            BytesIO object ของไฟล์ Word
        """
        doc = Document()
        
        # Header
        header = doc.add_heading("รายงานสรุปประวัติการวิเคราะห์", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_format = header.runs[0]
        header_format.font.color.rgb = RGBColor(46, 134, 222)
        
        # User info
        user_para = doc.add_paragraph()
        user_para.add_run(f"ผู้ใช้: ").bold = True
        user_para.add_run(username)
        
        date_para = doc.add_paragraph()
        date_para.add_run(f"วันที่สร้างรายงาน: ").bold = True
        date_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
        doc.add_paragraph()
        
        # Summary statistics
        doc.add_heading("📊 สถิติสรุป", level=2)
        
        summary_table = doc.add_table(rows=3, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ("จำนวนการวิเคราะห์ทั้งหมด", str(len(history_list))),
            ("วันที่ล่าสุด", history_list[0]['timestamp'] if history_list else "ไม่มี"),
            ("ไฟล์ที่วิเคราะห์", str(len(set([h['file_name'] for h in history_list]))))
        ]
        
        for idx, (key, value) in enumerate(summary_data):
            cells = summary_table.rows[idx].cells
            cells[0].text = key
            cells[1].text = value
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        # History list
        doc.add_heading("📜 รายการการวิเคราะห์", level=2)
        
        if history_list:
            history_table = doc.add_table(rows=len(history_list) + 1, cols=3)
            history_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = history_table.rows[0].cells
            header_cells[0].text = "ชื่อไฟล์"
            header_cells[1].text = "วันเวลา"
            header_cells[2].text = "ขนาด (ตัวอักษร)"
            
            for paragraph in header_cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            
            # Data rows
            for idx, entry in enumerate(history_list, start=1):
                cells = history_table.rows[idx].cells
                cells[0].text = entry['file_name']
                cells[1].text = entry['timestamp'][:19]
                cells[2].text = str(entry['file_size_chars'])
        else:
            doc.add_paragraph("ยังไม่มีประวัติการวิเคราะห์")
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer


# ตัวอย่างการใช้งาน
if __name__ == "__main__":
    print("Report Generator Module")
    print("-" * 50)
    
    # Test generate word report
    gen = ReportGenerator()
    
    sample_result = """
## 📊 ผลการวิเคราะห์ความสอดคล้อง

**1. วัตถุประสงค์ที่พบ:**
- เพื่อศึกษาความพึงพอใจของนักเรียน
- เพื่อปรับปรุงประสิทธิภาพการเรียน

**2. การตรวจสอบรายข้อ:**
- 🎯 ข้อ 1: ผ่าน เพราะสรุปผลมีข้อมูลเกี่ยวกับความพึงพอใจ
- 🎯 ข้อ 2: ผ่าน เพราะมีตัวเลขสถิติที่ชัดเจน

**3. ข้อแนะนำ:**
- โครงงานนี้มีความสอดคล้องดี
    """
    
    # Generate report
    doc_buffer = gen.generate_word_report(
        username="student1",
        file_name="โครงงาน_test.pdf",
        analysis_result=sample_result
    )
    
    # Save to file
    with open("รายงาน_test.docx", "wb") as f:
        f.write(doc_buffer.getvalue())
    
    print("✅ สร้างรายงาน Word สำเร็จ: รายงาน_test.docx")
    
    filename = gen.get_word_filename("โครงงาน_test.pdf", "student1")
    print(f"💾 ชื่อไฟล์: {filename}")
